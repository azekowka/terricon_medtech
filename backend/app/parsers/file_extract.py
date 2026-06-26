"""Multi-format (service_name, price) extraction for PDF / DOCX / XLSX (TZ 3.1).

Clinics that publish price lists as documents rather than HTML are handled here.
Each extractor returns list of (name, price_raw) pairs that flow through the same
normalization pipeline as the web scrapers.

Heavy libraries (pdfplumber/openpyxl/python-docx) are imported lazily so the API
starts fast and works even if a particular reader is absent.
"""
from __future__ import annotations

import re

from .html_extract import PRICE_RE

_NAME_OK = re.compile(r"[А-Яа-яЁё]{3,}")
_HEADER = re.compile(r"наименование|услуг|цена|прайс|прейскурант", re.IGNORECASE)


def _clean(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()


def _valid_name(s: str) -> bool:
    s = _clean(s)
    return len(s) >= 4 and bool(_NAME_OK.search(s)) and not _HEADER.search(s)


def _has_price(s: str) -> bool:
    return bool(PRICE_RE.search(s or "")) or bool(re.search(r"\d{3}", s or ""))


def _from_cells(name_cell: str, price_cell: str) -> tuple[str, str] | None:
    name, price = _clean(name_cell), _clean(price_cell)
    if _valid_name(name) and _has_price(price):
        return name, price
    return None


def extract_from_lines(lines: list[str]) -> list[tuple[str, str]]:
    """Parse 'Service name .... 1 200 ₸' style text lines."""
    pairs: list[tuple[str, str]] = []
    for line in lines:
        line = _clean(line)
        m = PRICE_RE.search(line)
        if not m:
            continue
        name = line[: m.start()].strip(" .—–-\t")
        if _valid_name(name):
            pairs.append((name, m.group(1)))
    return pairs


def extract_from_pdf(path: str) -> list[tuple[str, str]]:
    import pdfplumber

    pairs: list[tuple[str, str]] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            used_table = False
            for table in tables:
                for row in table:
                    if not row or len(row) < 2:
                        continue
                    pair = _from_cells(row[0] or "", row[-1] or "")
                    if pair:
                        pairs.append(pair)
                        used_table = True
            if not used_table:
                text = page.extract_text() or ""
                pairs += extract_from_lines(text.splitlines())
    return pairs


def extract_from_xlsx(path: str) -> list[tuple[str, str]]:
    from openpyxl import load_workbook

    pairs: list[tuple[str, str]] = []
    wb = load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            if not row or len(row) < 2:
                continue
            pair = _from_cells(str(row[0] or ""), str(row[-1] or ""))
            if pair:
                pairs.append(pair)
    wb.close()
    return pairs


def extract_from_docx(path: str) -> list[tuple[str, str]]:
    from docx import Document

    doc = Document(path)
    pairs: list[tuple[str, str]] = []
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            pair = _from_cells(cells[0].text, cells[-1].text)
            if pair:
                pairs.append(pair)
    # also scan paragraphs for inline "name price" lines
    pairs += extract_from_lines([p.text for p in doc.paragraphs])
    return pairs


def extract_file(path: str) -> list[tuple[str, str]]:
    """Dispatch on file extension. Returns de-duplicated (name, price_raw) pairs."""
    lower = path.lower()
    if lower.endswith(".pdf"):
        pairs = extract_from_pdf(path)
    elif lower.endswith((".xlsx", ".xlsm")):
        pairs = extract_from_xlsx(path)
    elif lower.endswith(".docx"):
        pairs = extract_from_docx(path)
    else:
        return []
    seen: set[str] = set()
    out: list[tuple[str, str]] = []
    for name, price in pairs:
        key = name.lower()
        if key not in seen:
            seen.add(key)
            out.append((name, price))
    return out
