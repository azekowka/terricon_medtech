"""Extraction of (service_name, price) from REAL clinic price lists.

Real price lists (data/real_prices/) are messy: multi-column tariff tables with
preambles, section headers and several price columns (resident / non-resident /
insurance). This module is column-aware for spreadsheets & DOCX tables (it finds
the "Наименование" and the citizen/resident "Цена" columns) and line-based for
PDFs (strip the leading tariff code, then take the first sane price = resident).

Supported: .xlsx .xlsm .xls .docx .pdf
"""
from __future__ import annotations

import re

from .base import parse_price

# --- price / name helpers --------------------------------------------------
_NUM = re.compile(r"\d[\d\s .,]*\d|\d")
_LEADING_CODE = re.compile(r"^\s*[A-Za-zА-Яа-я]?\d[\d.\-/]*\.?\s+")
_CYR = re.compile(r"[А-Яа-яЁё]")
_STARTS_OK = re.compile(r"^[A-ZА-ЯЁ0-9]")
MIN_PRICE, MAX_PRICE = 300, 1_500_000


def _clean(s) -> str:
    return re.sub(r"\s+", " ", str(s if s is not None else "").replace("\xa0", " ")).strip()


def _valid_name(s: str) -> bool:
    s = _clean(s)
    if not (5 <= len(s) <= 110) or len(_CYR.findall(s)) < 3:
        return False
    # must start with a capital letter or digit -> drops wrapped line fragments
    if not _STARTS_OK.match(s):
        return False
    low = s.lower()
    # skip column headers / section titles
    if any(h in low for h in ("наименование", "прейскурант", "приложение", "цена ", "стоимость,")):
        return False
    return True


def _prices_in(text, pick="first"):
    cands = []
    for m in _NUM.finditer(str(text or "")):
        v = parse_price(m.group(0))
        if v is not None and MIN_PRICE <= v <= MAX_PRICE:
            cands.append(v)
    if not cands:
        return None
    return cands[0] if pick == "first" else cands[-1]


# --- column-aware table extraction (xlsx/xls/docx tables) ------------------
def _find_columns(rows: list[list[str]]):
    """Locate the header row and (name_col, price_col)."""
    for r, row in enumerate(rows[:30]):
        cells = [_clean(c).lower() for c in row]
        name_col = next((i for i, c in enumerate(cells) if "наименование" in c), None)
        if name_col is None:
            name_col = next((i for i, c in enumerate(cells) if "услуг" in c and "ед" not in c), None)
        price_cols = [i for i, c in enumerate(cells)
                      if ("цена" in c or "стоимость" in c or "тенге" in c) and "без" not in c]
        if name_col is not None and price_cols:
            # prefer the resident/citizen column
            resident = [i for i in price_cols
                        if any(k in cells[i] for k in ("граждан", "резидент", "казахст", " рк", "(рк"))]
            price_col = resident[0] if resident else price_cols[0]
            return r, name_col, price_col
    return None, None, None


def _pairs_from_table(rows: list[list[str]]) -> list[tuple[str, str]]:
    header_r, name_col, price_col = _find_columns(rows)
    pairs: list[tuple[str, str]] = []
    if header_r is None:
        return pairs
    for row in rows[header_r + 1:]:
        if name_col >= len(row) or price_col >= len(row):
            continue
        name = _clean(row[name_col])
        if not _valid_name(name):
            continue
        price = _prices_in(row[price_col], pick="first")
        if price is None:
            # fall back: first sane price anywhere to the right of the name col
            price = _prices_in(" ".join(_clean(c) for c in row[name_col + 1:]), pick="first")
        if price is not None:
            pairs.append((name, str(price)))
    return pairs


# --- line-based extraction (PDF text, DOCX paragraphs) ---------------------
def _pair_from_line(line: str) -> tuple[str, str] | None:
    line = _clean(line)
    price = _prices_in(line, pick="first")
    if price is None:
        return None
    m = next((mm for mm in _NUM.finditer(line) if parse_price(mm.group(0)) == price), None)
    if not m:
        return None
    name = _LEADING_CODE.sub("", line[: m.start()]).strip(" .—–-\t")
    if _valid_name(name):
        return (name, str(price))
    return None


def _pairs_from_lines(lines) -> list[tuple[str, str]]:
    out = []
    for line in lines:
        p = _pair_from_line(line)
        if p:
            out.append(p)
    return out


# --- per-format loaders ----------------------------------------------------
def _xlsx_rows(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append([[c for c in row] for row in ws.iter_rows(values_only=True)])
    wb.close()
    return out


def _xls_rows(path):
    import xlrd
    wb = xlrd.open_workbook(path)
    out = []
    for sh in wb.sheets():
        out.append([[sh.cell_value(r, c) for c in range(sh.ncols)] for r in range(sh.nrows)])
    return out


def _docx_pairs(path):
    from docx import Document
    doc = Document(path)
    pairs = []
    for table in doc.tables:
        rows = [[c.text for c in row.cells] for row in table.rows]
        pairs += _pairs_from_table(rows)
    if not pairs:
        pairs += _pairs_from_lines([p.text for p in doc.paragraphs])
    return pairs


def _pdf_pairs(path):
    import pdfplumber
    pairs = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables() or []
            tpairs = []
            for table in tables:
                tpairs += _pairs_from_table([[c for c in row] for row in table])
            if tpairs:
                pairs += tpairs
            else:
                pairs += _pairs_from_lines((page.extract_text() or "").splitlines())
    return pairs


def extract_real_file(path: str) -> list[tuple[str, str]]:
    low = path.lower()
    if low.endswith((".xlsx", ".xlsm")):
        sheets = _xlsx_rows(path)
        pairs = [p for rows in sheets for p in _pairs_from_table(rows)]
    elif low.endswith(".xls"):
        sheets = _xls_rows(path)
        pairs = [p for rows in sheets for p in _pairs_from_table(rows)]
    elif low.endswith(".docx"):
        pairs = _docx_pairs(path)
    elif low.endswith(".pdf"):
        pairs = _pdf_pairs(path)
    else:
        return []
    # de-dup by normalized name (keep first)
    seen, out = set(), []
    for name, price in pairs:
        k = name.lower()
        if k not in seen:
            seen.add(k)
            out.append((name, price))
    return out
