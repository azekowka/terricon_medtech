"""Generates sample clinic price-list fixtures in PDF / XLSX / DOCX.

These represent the TZ source row "сайты городских и областных клиник | HTML,
PDF, DOCX" — public clinics that publish their price lists as documents rather
than HTML. The FixtureParser ingests them through the same normalization
pipeline as the web scrapers, demonstrating multi-format support (TZ 3.1).

Run (needs reportlab/openpyxl/python-docx):
    python scripts/build_fixtures.py
"""
import hashlib
import json
import os

from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet

# reportlab's built-in fonts have no Cyrillic glyphs (text extracts as garbage),
# so register a Unicode TTF. Fall back across common locations.
_FONT = "Helvetica"
for _cand in (
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\segoeui.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
):
    if os.path.exists(_cand):
        try:
            pdfmetrics.registerFont(TTFont("CyrFont", _cand))
            _FONT = "CyrFont"
            break
        except Exception:
            pass

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIX_DIR = os.path.join(ROOT, "data", "fixtures")

# Each fixture clinic: (filename, format, clinic meta dict, categories, factor)
CLINICS = [
    (
        "gorpoliklinika_astana.pdf", "pdf",
        {
            "clinic_name": "Городская поликлиника №1",
            "source": "gorpolik_astana", "city": "Астана",
            "address": "ул. Бейбітшілік, 12", "phone": "+7 7172 31 0001",
            "working_hours": "Пн–Пт 08:00–18:00, Сб 09:00–14:00",
            "website": "https://gp1.astana.kz", "rating": 4.2,
            "has_online_booking": False, "lat": 51.171, "lng": 71.428,
        },
        ["doctor", "laboratory", "procedure"], 0.82,
    ),
    (
        "obldiagcentr_shymkent.xlsx", "xlsx",
        {
            "clinic_name": "Областной диагностический центр",
            "source": "odc_shymkent", "city": "Шымкент",
            "address": "пр. Республики, 5", "phone": "+7 7252 40 0002",
            "working_hours": "Пн–Сб 08:00–17:00",
            "website": "https://odc.shymkent.kz", "rating": 4.3,
            "has_online_booking": True, "lat": 42.318, "lng": 69.601,
        },
        ["diagnostic", "laboratory"], 0.9,
    ),
    (
        "crb_aktobe.docx", "docx",
        {
            "clinic_name": "Центральная районная больница",
            "source": "crb_aktobe", "city": "Актобе",
            "address": "ул. Маресьева, 70", "phone": "+7 7132 22 0003",
            "working_hours": "Пн–Пт 08:00–17:00",
            "website": "https://crb.aktobe.kz", "rating": 4.0,
            "has_online_booking": False, "lat": 50.279, "lng": 57.207,
        },
        ["doctor", "laboratory", "diagnostic"], 0.8,
    ),
]


def _rng(*parts):
    h = hashlib.md5("|".join(parts).encode("utf-8")).hexdigest()
    return (int(h[:8], 16) % 1_000_000) / 1_000_000.0


def _rows(services_by_cat, cats, factor, clinic_name):
    rows = []
    for cat in cats:
        for svc in services_by_cat.get(cat, []):
            if _rng(clinic_name, svc["code"]) < 0.6:
                base = svc["base_price_kzt"] or 5000
                price = int(round(base * factor / 50.0) * 50)
                variants = [svc["name"], *svc["synonyms"]]
                idx = int(_rng(clinic_name, svc["code"], "v") * len(variants))
                name = variants[min(idx, len(variants) - 1)]
                rows.append((name, f"{price:,} ₸".replace(",", " ")))
    return rows


def write_pdf(path, clinic, rows):
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    for st in ("Title", "Normal", "Heading2"):
        styles[st].fontName = _FONT
    story = [
        Paragraph(clinic["clinic_name"], styles["Title"]),
        Paragraph(f"{clinic['city']}, {clinic['address']}", styles["Normal"]),
        Paragraph("Прейскурант на медицинские услуги (2025)", styles["Heading2"]),
        Spacer(1, 0.4 * cm),
    ]
    data = [["Наименование услуги", "Цена"]] + [list(r) for r in rows]
    table = Table(data, colWidths=[12 * cm, 4 * cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), _FONT),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F1F5F9")]),
    ]))
    story.append(table)
    doc.build(story)


def write_xlsx(path, clinic, rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Прайс"
    ws["A1"] = clinic["clinic_name"]
    ws["A2"] = f"{clinic['city']}, {clinic['address']}"
    ws.append([])
    ws.append(["Наименование услуги", "Цена, ₸"])
    for name, price in rows:
        ws.append([name, price])
    ws.column_dimensions["A"].width = 55
    ws.column_dimensions["B"].width = 16
    wb.save(path)


def write_docx(path, clinic, rows):
    doc = Document()
    doc.add_heading(clinic["clinic_name"], level=0)
    doc.add_paragraph(f"{clinic['city']}, {clinic['address']}")
    doc.add_heading("Прейскурант на медицинские услуги", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Наименование услуги", "Цена"
    for name, price in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = name, price
    doc.save(path)


def build():
    os.makedirs(FIX_DIR, exist_ok=True)
    with open(os.path.join(ROOT, "data", "services_dictionary.json"), encoding="utf-8") as f:
        services = json.load(f)
    by_cat = {}
    for s in services:
        by_cat.setdefault(s["category"], []).append(s)

    meta = {}
    for fname, fmt, clinic, cats, factor in CLINICS:
        rows = _rows(by_cat, cats, factor, clinic["clinic_name"])
        path = os.path.join(FIX_DIR, fname)
        if fmt == "pdf":
            write_pdf(path, clinic, rows)
        elif fmt == "xlsx":
            write_xlsx(path, clinic, rows)
        elif fmt == "docx":
            write_docx(path, clinic, rows)
        meta[fname] = {**clinic, "format": fmt, "rows": len(rows)}
        print(f"Wrote {fname}: {len(rows)} services ({fmt})")

    # clinic metadata sidecar consumed by FixtureParser
    with open(os.path.join(FIX_DIR, "_clinics.json"), "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f"Wrote metadata for {len(meta)} fixture clinics")


if __name__ == "__main__":
    build()
