"""Inspect the real clinic price files to understand their structure."""
import io
import os
import sys

sys.stdout.reconfigure(encoding="utf-8")
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIR = os.path.join(ROOT, "data", "real_prices")


def show_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    for ws in wb.worksheets[:2]:
        print(f"  sheet '{ws.title}' dims~{ws.max_row}x{ws.max_column}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 20:
                break
            vals = [str(c)[:30] if c is not None else "" for c in row]
            print("   ", " | ".join(vals))
    wb.close()


def show_xls(path):
    import xlrd
    wb = xlrd.open_workbook(path)
    for sh in wb.sheets()[:2]:
        print(f"  sheet '{sh.name}' {sh.nrows}x{sh.ncols}")
        for r in range(min(20, sh.nrows)):
            vals = [str(sh.cell_value(r, c))[:30] for c in range(sh.ncols)]
            print("   ", " | ".join(vals))


def show_docx(path):
    from docx import Document
    d = Document(path)
    print(f"  paragraphs={len(d.paragraphs)} tables={len(d.tables)}")
    for t in d.tables[:1]:
        print(f"  table rows={len(t.rows)} cols={len(t.columns)}")
        for r in t.rows[:12]:
            print("   ", " | ".join(c.text[:30] for c in r.cells))
    print("  -- paragraphs --")
    for p in [p.text for p in d.paragraphs if p.text.strip()][:12]:
        print("   ", p[:90])


def show_pdf(path):
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        print(f"  pages={len(pdf.pages)}")
        pg = pdf.pages[0]
        txt = pg.extract_text() or ""
        tables = pg.extract_tables() or []
        print(f"  page1: text_chars={len(txt)} tables={len(tables)} images={len(pg.images)}")
        if tables:
            for row in tables[0][:10]:
                print("   T|", " | ".join((str(c)[:28] if c else "") for c in row))
        print("  -- text head --")
        for line in txt.splitlines()[:18]:
            print("   ", line[:95])


HANDLERS = {".xlsx": show_xlsx, ".xls": show_xls, ".docx": show_docx, ".pdf": show_pdf}

target = sys.argv[1] if len(sys.argv) > 1 else None
for fname in sorted(os.listdir(DIR)):
    if target and target not in fname:
        continue
    ext = os.path.splitext(fname)[1].lower()
    h = HANDLERS.get(ext)
    print("=" * 80)
    print("FILE:", fname)
    if not h:
        print("  (no handler)")
        continue
    try:
        h(os.path.join(DIR, fname))
    except Exception as e:
        print(f"  ERROR: {type(e).__name__}: {e}")
